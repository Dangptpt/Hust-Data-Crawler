from docx import Document
import pandas as pd

# Đường dẫn tới file .docx
file_path = r'Quy-định-CLB-13_4_2020.docx'

# Mở file .docx
doc = Document(file_path)

# Mở file markdown để ghi
with open("Clean_data/QuyDinhCLB.txt", 'w', encoding='utf-8') as f:
    cnt = 0
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Nếu phần tử là đoạn văn
            paragraph = element.text
            if paragraph:  
                f.write(paragraph + '\n\n')  

        elif element.tag.endswith('tbl'):  # Nếu phần tử là bảng
            # Ghi tiêu đề bảng
            f.write("### Table:\n\n")

            # Lấy bảng tương ứng từ doc.tables
            table = doc.tables[cnt]  
            cnt += 1 
            # Chuyển đổi bảng thành DataFrame
            data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                data.append(row_data)
            
            # Chuyển đổi danh sách thành DataFrame
            df = pd.DataFrame(data)

            # Ghi DataFrame ở định dạng Markdown
            f.write(df.to_markdown(index=False))  
            f.write("\n\n" + "-" * 50 + "\n")  
print("Đã ghi nội dung ")
